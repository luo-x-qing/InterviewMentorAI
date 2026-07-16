"""
Word转Markdown转换器
支持.docx和.doc格式
"""
import os
import logging

logger = logging.getLogger(__name__)


def convert_docx_to_md(input_path: str, output_path: str) -> bool:
    """
    将Word文件转换为Markdown格式
    
    Args:
        input_path: Word文件路径
        output_path: 输出Markdown文件路径
        
    Returns:
        bool: 转换是否成功
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("请安装python-docx: pip install python-docx")
        return False
    
    try:
        doc = Document(input_path)
        content_parts = []
        
        for element in doc.element.body:
            # 处理段落
            if element.tag.endswith('}p'):
                paragraph = None
                for p in doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                
                if paragraph:
                    text = process_paragraph(paragraph)
                    if text:
                        content_parts.append(text)
            
            # 处理表格
            elif element.tag.endswith('}tbl'):
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table:
                    md_table = process_table(table)
                    if md_table:
                        content_parts.append(md_table)
        
        # 合并内容
        content = "\n\n".join(content_parts)
        
        # 清理内容
        content = clean_content(content)
        
        # 写入文件
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        logger.info(f"Word转换完成: {input_path} -> {output_path}")
        return True
        
    except Exception as e:
        logger.error(f"Word转换失败: {e}")
        return False


def process_paragraph(paragraph):
    """处理段落"""
    style_name = paragraph.style.name if paragraph.style else ""
    
    # 根据样式判断标题级别
    if style_name.startswith('Heading'):
        try:
            level = int(style_name.replace('Heading', '').strip())
            level = min(level, 6)  # Markdown最多6级标题
        except ValueError:
            level = 1
        return f"{'#' * level} {paragraph.text}"
    
    # 列表样式
    elif style_name.startswith('List'):
        return f"- {paragraph.text}"
    
    # 普通段落
    else:
        return paragraph.text


def process_table(table):
    """处理表格"""
    if not table.rows:
        return ""
    
    # 提取表头
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    
    # 提取数据行
    rows = []
    for row in table.rows[1:]:
        rows.append([cell.text.strip() for cell in row.cells])
    
    # 生成Markdown表格
    if not headers:
        return ""
    
    header_line = "| " + " | ".join(headers) + " |"
    separator = "| " + " | ".join(["---"] * len(headers)) + " |"
    data_lines = ["| " + " | ".join(row) + " |" for row in rows]
    
    return "\n".join([header_line, separator] + data_lines)


def clean_content(content: str) -> str:
    """清理转换后的内容"""
    import re
    
    # 移除多余空行
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    # 移除多余空格
    content = re.sub(r' +', ' ', content)
    
    # 清理特殊字符
    content = content.replace('\xa0', ' ')  # 非断行空格
    
    return content.strip()


if __name__ == '__main__':
    import sys
    if len(sys.argv) != 3:
        print("用法: python docx_to_md.py <输入DOCX> <输出MD>")
        sys.exit(1)
    
    convert_docx_to_md(sys.argv[1], sys.argv[2])
